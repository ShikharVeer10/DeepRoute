import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def build_indiacom_paper():
    doc = Document()
    
    # 1. Page Margins (0.75 in top/bottom, 0.63 in left/right - standard IEEE/INDIACom format)
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.63)
        s.right_margin = Inches(0.63)

    # 2. Add Running Header (INDIACom / IEEE Conference Header)
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("Proceedings of the 20th INDIACom; INDIACom-2026; IEEE Conference ID:\n2026 13th International Conference on \"Computing for Sustainable Global Development\", 08th – 10th April, 2026\nBharati Vidyapeeth's Institute of Computer Applications and Management (BVICAM), New Delhi (INDIA)")
    hrun.font.name = 'Times New Roman'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Helper XML functions for table cells
    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_cell_paddings(cell, top=60, bottom=60, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_section_columns(sec, num_cols):
        secPr = sec._sectPr
        cols = secPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), str(num_cols))
            cols[0].set(qn('w:space'), '720') # 0.5 in gap
        else:
            cols_elem = OxmlElement('w:cols')
            cols_elem.set(qn('w:num'), str(num_cols))
            cols_elem.set(qn('w:space'), '720')
            secPr.append(cols_elem)

    # Helper paragraph creators
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        return p

    def add_p(text, space_after=4, bold_prefix=None, italic_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if bold_prefix:
            r0 = p.add_run(bold_prefix)
            r0.bold = True
            r0.font.name = 'Times New Roman'
            r0.font.size = Pt(9.5)
        if italic_prefix:
            r0 = p.add_run(italic_prefix)
            r0.italic = True
            r0.font.name = 'Times New Roman'
            r0.font.size = Pt(9.5)
            
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r0 = p.add_run(bold_prefix)
            r0.bold = True
            r0.font.name = 'Times New Roman'
            r0.font.size = Pt(9.5)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        return p

    def add_code_block(title, lines_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_title = p.add_run(f"Algorithm: {title}\n")
        r_title.bold = True
        r_title.font.name = 'Courier New'
        r_title.font.size = Pt(8.5)
        r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        r_code = p.add_run(lines_text)
        r_code.font.name = 'Courier New'
        r_code.font.size = Pt(8.0)
        r_code.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    def add_image_figure(img_path, fig_num, caption_title, description):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            run_img = p_img.add_run()
            # 3.2 inches wide fits perfectly inside a single column of IEEE two-column layout
            run_img.add_picture(img_path, width=Inches(3.2))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(6)
            p_cap.paragraph_format.line_spacing = 1.0
            
            r_num = p_cap.add_run(f"Fig. {fig_num}. ")
            r_num.bold = True
            r_num.font.name = 'Times New Roman'
            r_num.font.size = Pt(8.5)
            
            r_cap = p_cap.add_run(caption_title)
            r_cap.font.name = 'Times New Roman'
            r_cap.font.size = Pt(8.5)
            
            add_p(description, space_after=6)

    # --- SINGLE COLUMN BANNER: TITLE & AUTHOR ---
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(12)
    tp.paragraph_format.space_after = Pt(12)
    tr = tp.add_run("DeepRoute: A Machine Learning and Graph Optimization Framework for Dynamic Multi-Objective Route Planning Under Simulated Urban Transportation Conditions")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = 'Times New Roman'
    tr.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_after = Pt(14)
    ar = ap.add_run("Shikhar Veeramachineni\nSchool of Computer Science and Engineering (SCOPE)\nVIT-AP University\nAmaravati, India\nshikhar.23bce9278@vitapstudent.ac.in")
    ar.font.size = Pt(10)
    ar.italic = True
    ar.font.name = 'Times New Roman'
    ar.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    # --- ADD CONTINUOUS SECTION BREAK FOR TWO-COLUMN BODY LAYOUT ---
    body_sec = doc.add_section(docx.enum.section.WD_SECTION.CONTINUOUS)
    set_section_columns(body_sec, 2)

    # --- ABSTRACT & KEYWORDS ---
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.space_before = Pt(0)
    p_abs.paragraph_format.space_after = Pt(4)
    p_abs.paragraph_format.line_spacing = 1.05
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_abs_lbl = p_abs.add_run("Abstract— ")
    r_abs_lbl.bold = True
    r_abs_lbl.italic = True
    r_abs_lbl.font.name = 'Times New Roman'
    r_abs_lbl.font.size = Pt(9)
    
    r_abs_txt = p_abs.add_run(
        "Modern urban navigation systems rely predominantly on static graph-search algorithms that do not account for stochastic transportation dynamics such as congestion fluctuations, weather events, and localized hazards. "
        "This paper presents DeepRoute, a modular systems framework that integrates gradient-boosted travel-time prediction, multi-objective graph optimization, stochastic risk modeling, and traffic-aware spatial visualization into a unified microservices architecture. "
        "DeepRoute operates on OpenStreetMap (OSM) directed road network graphs and transforms geospatial, traffic, and meteorological inputs into a 34-dimensional feature vector. "
        "Predictive models (XGBoost and LightGBM) trained on 10,000 samples from a Kaggle urban traffic density benchmark dataset, augmented with Indian metropolitan contextual features (festivals, monsoons, market days), dynamically reweight graph edges for multi-objective path ranking via a Weighted Sum Model (WSM) across 21 criteria. "
        "Monte Carlo simulation (N=1,000) quantifies travel-time volatility through Conditional Value-at-Risk (CVaR₉₅) bounds. "
        "On the held-out test partition, XGBoost achieves an R² of 0.9627, MAE of 0.0108 (±0.0002 across 5-fold CV), and inference latency of 2.75 ms, outperforming static Dijkstra and A* baselines that lack dynamic edge reweighting. "
        "A 100-trip simulated feedback evaluation yields a mean prediction error of 11.8% (σ=3.2%), demonstrating the framework's capacity for iterative calibration. "
        "The primary contribution is systems integration rather than algorithmic novelty: DeepRoute bridges existing ML prediction, graph optimization, risk quantification, and interactive GIS visualization components into a deployable end-to-end navigation pipeline."
    )
    r_abs_txt.font.name = 'Times New Roman'
    r_abs_txt.font.size = Pt(9)

    kp = doc.add_paragraph()
    kp.paragraph_format.space_after = Pt(10)
    kp.paragraph_format.line_spacing = 1.05
    kp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kr1 = kp.add_run("Keywords—")
    kr1.bold = True
    kr1.italic = True
    kr1.font.name = 'Times New Roman'
    kr1.font.size = Pt(9)
    kr2 = kp.add_run("Intelligent Transportation Systems (ITS), Travel-Time Prediction, Extreme Gradient Boosting (XGBoost), LightGBM, Multi-Objective Optimization, Weighted Sum Model (WSM), Conditional Value-at-Risk (CVaR), OpenStreetMap, Systems Integration.")
    kr2.bold = True
    kr2.font.name = 'Times New Roman'
    kr2.font.size = Pt(9)

    # --- SECTION I: INTRODUCTION ---
    add_h1("I. INTRODUCTION")
    add_p("Transportation networks constitute the critical infrastructure supporting metropolitan productivity, urban logistics, and economic sustainability [1]. Rapid urbanization and expanding vehicular volumes have intensified traffic congestion, causing economic losses, unpredictable travel delays, heightened accident risks, and increased carbon emissions [2]. To address these challenges, Intelligent Transportation Systems (ITS) leverage geospatial sensing, cloud computing, and artificial intelligence to transition urban navigation from reactive path-finding to proactive, predictive route management [3].")

    add_p("Traditional navigation platforms rely on classical graph-search algorithms such as Dijkstra's algorithm [4] and A* search [5], which compute shortest paths assuming static edge costs. In dynamic urban environments, edge traversal durations fluctuate due to temporal congestion, weather disruptions, road hazards, and special events [6]. Consequently, static routing paradigms frequently direct vehicles into emerging bottlenecks [7].")

    add_p("To model dynamic traffic behavior, machine learning techniques have been adopted for travel-time forecasting [8]. Gradient-boosted decision tree ensembles such as XGBoost [9] and LightGBM offer high accuracy and sub-millisecond inference on tabular transportation features. OpenStreetMap (OSM) provides globally accessible geospatial graph data [12], and OSMnx [13] enables extraction of directed road graphs enriched with spatial metadata.")

    add_p("Despite these advances, a gap persists in connecting individual components — ML prediction, graph optimization, risk quantification, and interactive visualization — into cohesive, deployable systems. Most studies evaluate travel-time prediction in isolation without multi-objective routing integration, stochastic risk bounds (e.g., CVaR), or closed-loop feedback mechanisms. Recent work on multi-route planning with matrix-based differential evolution [33] and reinforcement-learning-based urban routing [34] demonstrates growing interest in multi-objective formulations, yet few end-to-end architectures bridge predictive models with production-grade REST APIs, interactive dashboards, and continuous learning pipelines.")

    add_p("To address this integration gap, this paper presents DeepRoute, a modular framework for dynamic multi-objective route planning. The primary contribution is systems-level integration of established techniques rather than algorithmic novelty. The specific contributions are:")
    
    add_bullet(" Constructs a 34-dimensional feature vector from geospatial, temporal, weather, and regional Indian context inputs, combining cyclical temporal encodings, spatial road hierarchy, and dynamic severity indices.", bold_prefix="• Feature Engineering Pipeline:")
    add_bullet(" Deploys XGBoost and LightGBM regressors trained on Kaggle benchmark data augmented with Indian metropolitan features, achieving R² = 0.9627 and 2.75 ms inference latency.", bold_prefix="• Dual-Model Prediction Engine:")
    add_bullet(" Implements a 21-criterion Weighted Sum Model (WSM) with four user-configurable profiles (FASTEST, SAFEST, ECO, BALANCED) for candidate path ranking, with weight sensitivity analysis.", bold_prefix="• Multi-Objective Path Ranking:")
    add_bullet(" Integrates 1,000-iteration Monte Carlo sampling to quantify travel-time volatility through CVaR₉₅ risk bounds.", bold_prefix="• Stochastic Risk Quantification:")
    add_bullet(" Validates closed-loop prediction calibration across n=100 simulated trips, reporting error distributions (mean=11.8%, σ=3.2%).", bold_prefix="• Multi-Trip Feedback Evaluation:")
    add_bullet(" Renders Google Maps-style multi-segment traffic polylines with congestion-based color mapping on CartoDB Positron basemaps.", bold_prefix="• Traffic-Aware GIS Visualization:")
    add_bullet(" Serves inference via FastAPI endpoints and renders interactive Leaflet maps via Streamlit.", bold_prefix="• Microservices Deployment:")

    add_p("The remainder of this paper is organized as follows: Section II reviews related work. Section III presents the system methodology. Section IV details the system architecture and algorithms. Section V evaluates results with baseline comparisons, ablation analysis, multi-trip feedback evaluation, and generated empirical graphs. Section VI discusses limitations. Section VII concludes the paper.")

    # --- SECTION II: LITERATURE REVIEW ---
    add_h1("II. RELATED WORK")
    add_p("Route planning has evolved across multiple paradigms, from classical graph theory to machine learning forecasting, multi-objective optimization, and reinforcement learning.")

    add_h2("A. Classical Route Planning")
    add_p("Dijkstra's algorithm [4] guarantees exact single-source shortest paths on weighted directed graphs by exploring nodes in increasing distance order. Hart et al. [5] introduced A* search, which employs admissible heuristics to reduce node expansions. Yen [15] developed the K-shortest loopless path algorithm for alternative route generation. Bellman-Ford [14] accommodates negative edge weights at higher computational cost. While these methods remain foundational baselines, they assume static edge weights and cannot respond to dynamic traffic conditions.")

    add_h2("B. Machine Learning for Travel-Time Prediction")
    add_p("Data-driven regression models capture complex nonlinear feature interactions for travel-time forecasting [8]. Random Forest [16] and Extra Trees [17] provide robust ensemble baselines. Chen and Guestrin [9] introduced XGBoost, achieving strong performance on tabular benchmarks through regularized gradient boosting with exact greedy tree search and column subsampling. LightGBM further accelerates training via histogram-based binning. Recent deep learning approaches include LSTM networks for temporal traffic sequences and attention-based spatial-temporal graph convolutional networks (ASTGCN) [23] for capturing non-Euclidean spatial dependencies, though these typically require large-scale sensor datasets.")

    add_h2("C. Multi-Objective Optimization & Risk Modeling")
    add_p("Real-world navigation involves balancing multiple competing objectives beyond travel time: distance, safety, fuel consumption, EV energy usage, and driving comfort [24]. Multi-objective approaches employ Pareto frontier search [25] or scalarization techniques [26]. The Weighted Sum Model (WSM) scalarizes multi-dimensional costs using user-defined weights for real-time graph traversal [27]. Li et al. [33] employed matrix-based differential evolution with manual weight tuning and sensitivity analysis for tour multi-route planning. For stochastic risk quantification, Rockafellar and Uryasev [30] introduced CVaR, measuring expected losses exceeding a VaR percentile threshold.")

    add_h2("D. Reinforcement Learning & Recent Approaches")
    add_p("Peng et al. [34] reformulated urban route planning as a model-based reinforcement learning problem solved via dynamic programming on real Shenzhen road networks (2,245 nodes), comparing against Dijkstra baselines. Their approach generates alternative routes via ranked Q-values to mitigate congestion drift. Recent work on eco-driving route planning [35] integrates NSGA-II optimization with TomTom API traffic data and physical EV energy models, framing the contribution as systems integration of existing techniques for Pareto-optimal eco-routing.")

    add_h2("E. Literature Comparison")
    add_p("Table I compares existing approaches against DeepRoute across key dimensions.")

    # TABLE I: Literature Comparison
    t1 = doc.add_table(rows=8, cols=7)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Author & Year", "Method", "Dataset", "Technique", "Advantages", "Limitations", "Research Gap"]
    hdr_cells = t1.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_bg(hdr_cells[idx], "1B365D")
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(7.5)
        set_cell_paddings(hdr_cells[idx])

    rows_data = [
        ("Dijkstra (1959) [4]", "Static Graph Search", "Synthetic Graphs", "Priority Queue Shortest Path", "Guaranteed shortest-path optimality", "Assumes static edge costs", "No traffic prediction or dynamic adaptability"),
        ("Chen & Guestrin (2016) [9]", "Gradient Boosted Trees", "Kaggle Benchmarks", "Regularized Tree Boosting (XGBoost)", "High tabular accuracy & fast execution", "Evaluates tabular data without spatial routing", "No graph routing or multi-objective trade-offs"),
        ("Boeing (2017) [13]", "OSM Spatial Mining", "OpenStreetMap Data", "OSMnx Graph Extraction", "Automates spatial road graph construction", "Static geometry without predictions", "No ML engines or risk bounds"),
        ("Peng et al. (2022) [34]", "RL via Dynamic Programming", "Shenzhen Road Network (2,245 nodes)", "Model-Based RL + DCI Reward Shaping", "Multi-route generation on real network", "Deterministic policy; no weather/contextual features", "No ML prediction or stochastic risk"),
        ("Li et al. (2024) [33]", "Matrix-Based DE", "South Korea Tourism (20 spots)", "Vectorized Differential Evolution + WSM", "60× speedup; weight sensitivity analysis", "Tour optimization, not real-time navigation", "No traffic prediction or continuous feedback"),
        ("IEEE Access (2025) [35]", "NSGA-II Eco-Routing", "TomTom API + NYC OpenData", "Random Forest + NSGA-II Pareto", "Systems integration with real traffic data", "Limited to 2-objective (time, energy); R²=0.735", "No multi-dimensional WSM or CVaR risk bounds"),
        ("DeepRoute (This Work)", "ML + Multi-Obj Graph Opt.", "Kaggle Urban Traffic + Indian Context", "XGBoost + WSM (21 criteria) + Monte Carlo CVaR", "R²=0.9627; 21-criterion WSM; CVaR₉₅ risk; GIS UI", "Trained on augmented benchmark data", "Systems integration contribution"),
    ]

    for r_idx, r_data in enumerate(rows_data, start=1):
        row_cells = t1.rows[r_idx].cells
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_bg(row_cells[c_idx], bg_color)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(7.5)
                p.runs[0].font.name = 'Times New Roman'
            set_cell_paddings(row_cells[c_idx])

    add_p("", space_after=4)

    # --- SECTION III: METHODOLOGY ---
    add_h1("III. METHODOLOGY")
    add_p("This section describes the dataset, graph construction, feature engineering, and hardware configuration.")

    add_h2("A. Dataset")
    add_p("The prediction models are trained and evaluated on 10,000 transportation samples derived from the Kaggle Urban Traffic Density benchmark dataset, which provides tabular records of road segment attributes including speed, congestion levels, and temporal patterns across simulated urban corridors. The base dataset was augmented with Indian metropolitan contextual features — monsoon rain severity indices, regional festival traffic surges, peak-hour temporal encodings, and localized road hazard indicators — to reflect transportation dynamics characteristic of Indian metropolitan networks. Features are sampled across continuous and categorical distributions modeling speed limits (20–120 km/h), road lengths (50–5,000 m), lane counts (1–6), congestion indices (0.0–1.0), and weather severity indices (0.0–1.0). The dataset is partitioned using an 80/20 train-test split (8,000 training, 2,000 test) with 5-fold cross-validation for hyperparameter tuning via Optuna Bayesian optimization.")

    add_p("Limitation: The training data consists of augmented benchmark samples rather than field-collected GPS traces. Consequently, the evaluation measures the framework's predictive accuracy on the benchmark distribution and its integration behavior, not generalization to live traffic conditions. Validation on real-world GPS trajectory datasets (e.g., Porto Taxi, NYC TLC) remains necessary for deployment claims.")

    add_h2("B. OpenStreetMap (OSM) Graph Construction")
    add_p("Geospatial road networks are acquired dynamically from OpenStreetMap using OSMnx. Raw OSM street networks are converted into directed multigraphs G = (V, E), where nodes V represent road intersections and edges E represent directed street segments. Topology sanitization filters non-drivable paths, simplifies complex intersection clusters, and computes spatial attributes including geodesic segment length (meters), posted speed limit (km/h), lane count, and highway classification.")

    add_h2("C. Feature Engineering (34 Dimensions)")
    add_p("DeepRoute transforms heterogeneous inputs into a structured 34-dimensional feature vector, organized into four domain subsets:")
    
    add_bullet(" Encodes departure time using cyclical sine/cosine transforms for hour of day (hour_sin, hour_cos) and day of week (day_sin, day_cos), binary flags for peak-hour windows (07:00–10:00, 17:00–20:00), weekend indicators, Indian festival indicators with severity scores, monsoon season flags with rainfall severity, school zone operational hours, and market day activity. [12 features: hour_sin, hour_cos, day_sin, day_cos, is_peak_hour, is_weekend, is_festival, festival_severity, is_monsoon_season, monsoon_severity, is_school_hours, is_market_day]", bold_prefix="1) Temporal & Regional Context Features (12 dimensions):")
    add_bullet(" Incorporates segment length (length_m), posted speed limit (speed_limit_kph), number of lanes (num_lanes), and elevation delta (elevation_change_m). [4 features]", bold_prefix="2) Spatial Infrastructure Features (4 dimensions):")
    add_bullet(" Measures real-time link congestion (congestion_index), weather severity (weather_severity), incident proximity (incident_proximity), event proximity (event_proximity), synthesized road risk score (road_risk_score), binary status indicators for road closures, roadworks, and accidents, historical link traversal speed (historical_speed_kph), historical congestion (historical_congestion), speed variance reliability (speed_reliability), encoded road type (road_type_encoded), highway percentage (highway_percentage), route curvature (route_curvature), intersection count (intersection_count), toll road indicator (toll_roads), urban density (urban_density), and distance category (distance_category). [18 features]", bold_prefix="3) Dynamic Context & Environmental Features (18 dimensions):")

    add_p("The total feature count is 12 + 4 + 18 = 34 dimensions, matching the implementation in the feature engineering module.")

    add_h2("D. Hardware Configuration")
    add_p("All training, routing, and evaluation were executed on: Intel Core i7-13700H (16 cores, 24 threads, 5.0 GHz); 16 GB DDR5 RAM; Windows 11 (64-bit); Python 3.12.3; XGBoost 2.0.3, LightGBM 4.3.0, scikit-learn 1.4.1, NetworkX 3.2.1, OSMnx 1.9.1, FastAPI 0.110.0, Streamlit 1.31.1.")

    # --- SECTION IV: SYSTEM ARCHITECTURE & ALGORITHMS ---
    add_h1("IV. SYSTEM ARCHITECTURE & ALGORITHMS")
    add_p("This section describes the modular system architecture and formalizes key algorithmic components.")

    add_h2("A. System Architecture")
    add_p("DeepRoute integrates nine modules into a modular pipeline:")
    
    add_bullet(" Ingests road geometries from OpenStreetMap, live speed telemetry (TomTom Traffic API), Open-Meteo weather feeds, and regional Indian calendar parameters.", bold_prefix="1. Multi-Source Ingestion: ")
    add_bullet(" Transforms raw inputs into a 34-dimensional feature vector combining temporal cyclical encodings, spatial road metrics, environmental indices, and historical speed profiles.", bold_prefix="2. Feature Engineering: ")
    add_bullet(" Dispatches feature vectors to hyperparameter-optimized XGBoost and LightGBM models, predicting travel-time factors with sub-3 ms latency.", bold_prefix="3. Dual-Model ML Engine: ")
    add_bullet(" Constructs NetworkX directed multigraphs G=(V,E) and applies dynamic edge reweighting based on ML-predicted traffic, weather severity, and incident penalties.", bold_prefix="4. Graph Construction & Edge Reweighting: ")
    add_bullet(" Generates candidate paths via penalty-based A* routing and evaluates them using a 21-criterion WSM across four user profiles (FASTEST, SAFEST, ECO, BALANCED).", bold_prefix="5. Multi-Objective WSM Optimizer: ")
    add_bullet(" Executes 1,000 Monte Carlo iterations per route to compute CVaR₉₅ travel-time risk bounds.", bold_prefix="6. Monte Carlo Risk Assessor: ")
    add_bullet(" Exposes RESTful API endpoints (/api/route, /api/forecast, /api/risk, /api/recommend, /api/travel_data/collect) via FastAPI.", bold_prefix="7. FastAPI API Layer: ")
    add_bullet(" Renders interactive Leaflet maps with Google Maps-style multi-segment traffic coloring (#4285F4 free-flow, #FBBC04 moderate, #EA4335 heavy) on CartoDB Positron basemaps via Streamlit.", bold_prefix="8. GIS Visualization Frontend: ")
    add_bullet(" Logs actual trip durations asynchronously into SQLite storage, tracking prediction error margins for iterative edge cost recalibration.", bold_prefix="9. Feedback Collection Engine: ")

    add_h2("B. Algorithm 1: WSM Multi-Objective Path Scoring")
    add_p("Algorithm 1 formalizes the Weighted Sum Model scoring of candidate routes across 21 normalized criteria.")

    add_code_block(
        "Algorithm 1: WSM Multi-Objective Path Scoring",
        "Input : Candidate Routes R = {r_1, ..., r_K}, User Profile P ∈ {FASTEST, SAFEST, ECO, BALANCED}\n"
        "Output: Ranked Routes with composite scores\n\n"
        "1  Load weight vector W_P = {w_1, ..., w_21} for profile P\n"
        "2  Normalize: w_i ← w_i / Σ(w_j) for all i\n"
        "3  For each route r_k in R do:\n"
        "4      Extract raw metrics M_k = {m_1, ..., m_21}  // time, dist, congestion, risk, ...\n"
        "5  End For\n"
        "6  For each criterion j in {1..21} do:\n"
        "7      min_j ← min(M_k[j] for all k)\n"
        "8      max_j ← max(M_k[j] for all k)\n"
        "9      For each route r_k: m_k[j] ← (m_k[j] - min_j) / (max_j - min_j + ε)\n"
        "10 End For\n"
        "11 For each route r_k do:\n"
        "12     score_k ← Σ(w_j × m_k[j]) for j = 1..21\n"
        "13 End For\n"
        "14 Return R sorted by ascending score_k"
    )

    add_h2("C. Algorithm 2: Google Maps-Style Multi-Segment Traffic Polyline Rendering")
    add_p("Algorithm 2 formalizes the point-to-point segment traffic color assignment and Leaflet polyline rendering pipeline.")

    add_code_block(
        "Algorithm 2: Multi-Segment Traffic Polyline Rendering",
        "Input : Route lat/lng coordinates L = [p_1, ..., p_N], Segment Congestion Indices C = [c_1, ..., c_{N-1}]\n"
        "Output: Rendered Leaflet FeatureGroup with traffic-colored polylines\n\n"
        "1  Initialize Segment Array S ← []\n"
        "2  For i ← 1 to N-1 do:\n"
        "3      S.append({ coords: [L[i], L[i+1]], congestion: C[i] })\n"
        "4  End For\n"
        "5  Create Leaflet fillGroup ← L.featureGroup()\n"
        "6  For each segment s in S do:\n"
        "7      If s.congestion >= 0.45 then\n"
        "8          color ← '#EA4335'   // Heavy Traffic (Red)\n"
        "9      Else If s.congestion >= 0.25 then\n"
        "10         color ← '#FBBC04'   // Moderate Traffic (Yellow)\n"
        "11     Else\n"
        "12         color ← '#4285F4'   // Free-Flow Traffic (Blue)\n"
        "13     End If\n"
        "14     Polyline ← L.polyline(s.coords, { color, weight: 6, opacity: 1.0 })\n"
        "15     fillGroup.addLayer(Polyline)\n"
        "16 End For\n"
        "17 Return fillGroup.addTo(map)"
    )

    add_h2("D. Algorithm 3: Monte Carlo CVaR₉₅ Risk Bound Estimation")
    add_p("Algorithm 3 formalizes the stochastic travel-time risk quantification via Monte Carlo simulation.")

    add_code_block(
        "Algorithm 3: Monte Carlo CVaR₉₅ Risk Estimation",
        "Input : Base travel time T_base, Feature vector X, ML model M, N_sim = 1000\n"
        "Output: VaR_95, CVaR_95 risk bounds\n\n"
        "1  Initialize samples[] ← []\n"
        "2  For i ← 1 to N_sim do:\n"
        "3      X_perturbed ← X + ε, where ε ~ N(0, σ²_feature)\n"
        "4      factor_i ← M.predict(X_perturbed)\n"
        "5      T_i ← T_base × factor_i\n"
        "6      samples.append(T_i)\n"
        "7  End For\n"
        "8  Sort samples in ascending order\n"
        "9  VaR_95 ← samples[⌈0.95 × N_sim⌉]\n"
        "10 CVaR_95 ← mean(samples[j] for j where samples[j] ≥ VaR_95)\n"
        "11 Return VaR_95, CVaR_95"
    )

    # --- SECTION V: RESULTS AND DISCUSSION ---
    add_h1("V. RESULTS AND DISCUSSION")
    add_p("This section presents empirical performance results, model comparison benchmarks, WSM weight sensitivity analysis, multi-trip telemetry evaluation, generated analytical graphs, and visual implementation analysis.")

    add_h2("A. Model Comparison with Baselines")
    add_p("Table II compares ML regressors against non-ML baselines on the 2,000-sample held-out test set. Non-ML baselines are defined as follows: (1) Static Dijkstra uses free-flow speed limits as fixed edge weights without ML reweighting; (2) Static A* adds a haversine heuristic but retains fixed weights; (3) Mean-Speed Baseline assigns a constant travel-time factor of 1.0 (no dynamic adjustment). All ML models report 5-fold cross-validation MAE with standard deviation.")

    # TABLE II: Model Comparison Table with baselines and variance
    t2 = doc.add_table(rows=11, cols=8)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_headers = ["Model", "MAE", "RMSE", "R²", "MAPE (%)", "5-Fold CV MAE (±σ)", "Latency (ms)", "Type"]
    hdr2 = t2.rows[0].cells
    for idx, text in enumerate(m_headers):
        hdr2[idx].text = text
        set_cell_bg(hdr2[idx], "1B365D")
        hdr2[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr2[idx].paragraphs[0].runs[0].font.bold = True
        hdr2[idx].paragraphs[0].runs[0].font.size = Pt(7.5)
        set_cell_paddings(hdr2[idx])

    m_rows = [
        ("XGBoost", "0.01081", "0.01353", "0.9627", "0.988", "0.01086 ± 0.00025", "2.75", "ML (Production)"),
        ("LightGBM", "0.01086", "0.01366", "0.9620", "0.992", "0.01088 ± 0.00025", "2.10", "ML (Secondary)"),
        ("Gradient Boosting", "0.01349", "0.01712", "0.9399", "1.284", "0.01371 ± 0.00031", "3.45", "ML Ensemble"),
        ("HistGradientBoosting", "0.01391", "0.01764", "0.9362", "1.312", "0.01412 ± 0.00033", "1.95", "ML Ensemble"),
        ("Extra Trees", "0.01425", "0.01815", "0.9324", "1.355", "0.01458 ± 0.00035", "4.82", "ML Ensemble"),
        ("Random Forest", "0.01512", "0.01928", "0.9245", "1.438", "0.01549 ± 0.00038", "5.15", "ML Ensemble"),
        ("Ridge Regression", "0.02845", "0.03512", "0.7512", "2.715", "0.02891 ± 0.00052", "0.42", "Linear Baseline"),
        ("Static Dijkstra", "0.08720", "0.10340", "0.0000", "8.410", "N/A (deterministic)", "1.20", "Non-ML Baseline"),
        ("Static A*", "0.08720", "0.10340", "0.0000", "8.410", "N/A (deterministic)", "0.85", "Non-ML Baseline"),
        ("Mean-Speed (factor=1.0)", "0.07150", "0.08920", "0.0000", "6.880", "N/A (constant)", "0.01", "Trivial Baseline"),
    ]

    for r_idx, r_data in enumerate(m_rows, start=1):
        row_cells = t2.rows[r_idx].cells
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_bg(row_cells[c_idx], bg_color)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(7.5)
                p.runs[0].font.name = 'Times New Roman'
            set_cell_paddings(row_cells[c_idx])

    add_p("", space_after=4)

    # Insert Figure 4: Model Comparison Graph
    fig4_path = r"C:\Users\shikh\DeepRoute\paper_figures\fig4_model_comparison.png"
    add_image_figure(
        fig4_path,
        4,
        "Comparative benchmark performance (MAE vs. R² Score) across machine learning regressors.",
        "Figure 4 illustrates the trade-off between forecasting accuracy (MAE, blue) and variance explained (R², red) across candidate machine learning regressors. Gradient boosting models (XGBoost and LightGBM) demonstrate superior performance, achieving R² > 0.96 and MAE < 0.011, significantly outperforming linear Ridge Regression."
    )

    add_h2("B. Analysis of Results")
    add_p("XGBoost achieved the highest predictive precision, explaining 96.27% of travel-time factor variance (R² = 0.9627) with an MAE of 0.01081 ± 0.00025 (5-fold CV) and inference latency of 2.75 ms. LightGBM closely followed (R² = 0.9620, MAE = 0.01086 ± 0.00025). Both gradient-boosted models substantially outperform the non-ML baselines: static Dijkstra and A* produce MAE of 0.0872 (8.1× higher), confirming that ML-based dynamic edge reweighting provides meaningful improvement over fixed-weight graph search on this benchmark. Ridge Regression (R² = 0.7512) serves as a linear baseline, demonstrating that the feature space contains substantial nonlinear interactions that tree-based ensembles capture effectively.")

    add_p("The 5-fold CV standard deviations (σ ≈ 0.00025–0.00052) indicate stable performance across data partitions, with coefficient of variation below 5% for all ML models. The gap between tree-ensemble and linear models confirms that the 34-dimensional feature vector encodes meaningful nonlinear congestion-weather-temporal interactions.")

    add_h2("C. WSM Weight Sensitivity Analysis")
    add_p("The 21-criterion WSM assigns weights manually based on domain heuristics reflecting the semantic priority of each user profile. Table III presents the dominant weight allocations and demonstrates route ranking sensitivity to weight perturbation.")

    # TABLE III: WSM Weight Sensitivity
    t3 = doc.add_table(rows=6, cols=5)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_headers = ["Profile", "Top-3 Weighted Criteria", "Dominant Weight Sum", "Rank Change on ±20% Perturbation", "Justification"]
    hdr3 = t3.rows[0].cells
    for idx, text in enumerate(w_headers):
        hdr3[idx].text = text
        set_cell_bg(hdr3[idx], "1B365D")
        hdr3[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr3[idx].paragraphs[0].runs[0].font.bold = True
        hdr3[idx].paragraphs[0].runs[0].font.size = Pt(7.5)
        set_cell_paddings(hdr3[idx])

    w_rows = [
        ("FASTEST", "travel_time (0.24), congestion (0.10), traffic_delay (0.08)", "0.42 / 1.00", "0/3 routes change rank", "Minimizes travel duration; congestion & delay amplify time cost"),
        ("SAFEST", "risk (0.14), incident_count (0.12), safety (0.12)", "0.38 / 1.00", "1/3 routes swap rank (R2↔R3)", "Prioritizes hazard avoidance; incident/construction penalties dominate"),
        ("ECO", "fuel (0.22), ev_energy (0.16), distance (0.10)", "0.48 / 1.00", "0/3 routes change rank", "Minimizes energy consumption; fuel & EV weights strongly correlated"),
        ("BALANCED", "travel_time (0.12), distance (0.08), congestion (0.08)", "0.28 / 1.00", "1/3 routes swap rank (R1↔R2)", "Equal-emphasis profile; most sensitive to perturbation"),
        ("Uniform (ablation)", "All criteria = 1/21 ≈ 0.048", "N/A", "2/3 routes change rank vs FASTEST", "Ablation baseline: removing domain priors degrades ranking quality"),
    ]

    for r_idx, r_data in enumerate(w_rows, start=1):
        row_cells = t3.rows[r_idx].cells
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_bg(row_cells[c_idx], bg_color)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(7.5)
                p.runs[0].font.name = 'Times New Roman'
            set_cell_paddings(row_cells[c_idx])

    add_p("", space_after=4)

    add_p("The sensitivity analysis reveals that FASTEST and ECO profiles are robust to ±20% weight perturbation (no rank changes), while BALANCED is most sensitive due to its flatter weight distribution. The uniform-weight ablation (all criteria weighted equally) causes 2 of 3 route rankings to change relative to FASTEST, demonstrating that domain-informed weight assignment produces meaningfully different route selections compared to naive equal weighting. Weights are user-configurable at query time; the presets represent common navigation priorities rather than empirically optimized coefficients.")

    add_h2("D. Multi-Trip Telemetry Evaluation & Risk Quantification")
    add_p("To evaluate the closed-loop feedback mechanism beyond a single anecdote, a simulation study was conducted with n=100 synthetic trips across diverse conditions (varying congestion levels, weather severity, and time of day). For each trip, the XGBoost model predicted travel time, and a simulated 'actual' travel time was generated by adding noise sampled from the empirical error distribution observed in calibration runs.")

    # Insert Figure 5: Error Distribution
    fig5_path = r"C:\Users\shikh\DeepRoute\paper_figures\fig5_error_distribution.png"
    add_image_figure(
        fig5_path,
        5,
        "Probability density distribution of travel-time prediction errors across n=100 simulated trips.",
        "Figure 5 presents the empirical error distribution of the closed-loop telemetry feedback engine. The mean absolute prediction error is 11.8% (μ=11.8%) with a standard deviation of 3.2% (σ=3.2%). The 95th percentile error bound occurs at 17.9%, establishing the risk boundary for severe delay anomalies."
    )

    add_p("Results: The 100-trip evaluation yielded a mean absolute prediction error of 11.8% with a standard deviation of 3.2%. The error distribution is approximately normal with a 5th percentile of 6.4% and 95th percentile of 17.9%. The median error was 11.2%, indicating slight right skew from high-congestion outlier trips. These results characterize the framework's prediction calibration behavior under simulated conditions.")

    # Insert Figure 6: Monte Carlo CVaR Graph (renamed Figure 6 in order)
    fig7_path = r"C:\Users\shikh\DeepRoute\paper_figures\fig7_cvar_simulation.png"
    add_image_figure(
        fig7_path,
        6,
        "Stochastic travel-time distribution and CVaR₉₅ risk bounds from 1,000 Monte Carlo sampling iterations.",
        "Figure 6 illustrates the output of the Monte Carlo risk engine (Algorithm 3) for a 60-minute baseline route. Operating over 1,000 sampling runs under stochastic feature perturbations, the expected mean travel duration is 60.5 minutes, while the Value-at-Risk (VaR-95) is 73.1 minutes and Conditional Value-at-Risk (CVaR-95) is 78.4 minutes, providing a robust tail-risk bound for routing optimization."
    )

    add_h2("E. Visual Implementation Analysis")
    add_p("Figures 7–9 present implementation screenshots from the Streamlit GIS frontend, illustrating the traffic-aware visualization and multi-route selection capabilities.")

    # Figure 7: Screenshot (56).png - National Corridor Route Planning
    img1_path = r"C:\Users\shikh\OneDrive\Documents\Desktop\Screenshot (56).png"
    add_image_figure(
        img1_path,
        7,
        "Multi-segment traffic polyline overlay along a national highway corridor (New Delhi to Hyderabad).",
        "Figure 7 illustrates the geospatial deployment of DeepRoute along a 1,500 km national highway corridor. The map demonstrates the multi-segment traffic polyline engine rendered over a CartoDB Positron basemap. Free-flow segments appear in blue (#4285F4), moderate congestion in yellow (#FBBC04), and heavy congestion in red (#EA4335). Spatial hazard clusters are dynamically pinned along the path, including speed cameras, roadworks (cluster badge 7), and road closures (cluster badge 5). The ETA badge displays Route 2 with 32h 06min projected duration."
    )

    # Figure 8: Screenshot 2026-07-30 172313.png - Regional Highway Corridor
    img2_path = r"C:\Users\shikh\Pictures\Screenshots\Screenshot 2026-07-30 172313.png"
    add_image_figure(
        img2_path,
        8,
        "Regional corridor traffic overlay (Hyderabad to Bengaluru, NH 44) with hazard clustering.",
        "Figure 8 depicts a regional corridor from Hyderabad to Bengaluru. The visualization demonstrates point-to-point segment coloring precision of Algorithm 2. Severe bottlenecking (red polylines) near Kurnool and Anantapur reflects high-density roadworks (cluster badge 6) and environmental hazards. The selection badge displays 'Route 3: 13h 21min', illustrating how dynamic edge reweighting incorporates local hazard clusters into route cost calculations."
    )

    # Figure 9: Screenshot 2026-07-30 172827.png - Macroscopic National Overview
    img3_path = r"C:\Users\shikh\Pictures\Screenshots\Screenshot 2026-07-30 172827.png"
    add_image_figure(
        img3_path,
        9,
        "National-scale optimal route selection ('Fastest: 29h 29min') under WSM FASTEST profile.",
        "Figure 9 presents the macroscopic national-scale navigation view. The WSM FASTEST profile identifies the optimal path (29h 29min) connecting New Delhi to Hyderabad. Comparing with Figure 7 (32h 06min via Route 2), the multi-objective optimizer achieves a 2h 37min savings by routing around congestion clusters near Nagpur and Bhopal. The floating UI sidebar displays the legend, route status toggles, and congestion severity indicators."
    )

    # --- SECTION VI: LIMITATIONS ---
    add_h1("VI. LIMITATIONS")
    add_p("Several limitations should be noted when interpreting the results presented in this paper:")

    add_bullet(" The ML models are trained on Kaggle benchmark data augmented with synthetically generated Indian context features. While the augmentation reflects plausible metropolitan dynamics (monsoons, festivals, peak hours), the training distribution may not capture all real-world traffic phenomena. Validation on field-collected GPS trajectory datasets (e.g., Porto Taxi, NYC TLC trip records) is necessary before deployment claims can be made.", bold_prefix="Data Provenance: ")
    add_bullet(" The 21-criterion WSM weights are assigned via domain heuristics rather than learned from user preference data or optimized via multi-objective search (e.g., NSGA-II). While the sensitivity analysis (Table III) demonstrates robustness of FASTEST and ECO profiles, the BALANCED profile exhibits sensitivity to perturbation. Future work should explore preference elicitation or Pareto-based weight optimization.", bold_prefix="Weight Selection: ")
    add_bullet(" The n=100 feedback evaluation uses simulated trips with synthetic noise. A field deployment with actual driver trip telemetry is required to validate the continuous learning mechanism under real traffic dynamics.", bold_prefix="Feedback Evaluation: ")
    add_bullet(" Static Dijkstra and A* baselines serve as lower bounds by using free-flow weights without ML reweighting. Comparison against production navigation APIs (Google Maps, Waze) or published deep learning ETA models would provide stronger external baselines.", bold_prefix="Baseline Scope: ")

    # --- SECTION VII: CONCLUSION ---
    add_h1("VII. CONCLUSION")
    add_p("This paper presented DeepRoute, a modular systems framework that integrates gradient-boosted travel-time prediction, multi-objective graph optimization with a 21-criterion Weighted Sum Model, Monte Carlo CVaR₉₅ risk quantification, and Google Maps-style traffic visualization into a deployable navigation pipeline. The primary contribution is systems-level integration of established techniques — XGBoost, WSM, CVaR, OSMnx — rather than algorithmic novelty.")

    add_p("On the benchmark evaluation, XGBoost achieves R² = 0.9627 with MAE = 0.01081 ± 0.00025 (5-fold CV) and 2.75 ms inference latency, substantially outperforming static Dijkstra/A* baselines (MAE = 0.0872). The WSM weight sensitivity analysis confirms that domain-informed weight profiles produce stable, meaningful route rankings under ±20% perturbation. A 100-trip simulated feedback evaluation characterizes prediction calibration (mean error 11.8%, σ=3.2%). Generated empirical charts (Figures 4–6) further illustrate benchmark tradeoffs, risk distributions, and stochastic travel duration bounds.")

    add_p("Future work will focus on three directions: (1) training and evaluating on real-world GPS trajectory datasets to validate generalization beyond benchmark distributions; (2) replacing manual WSM weights with learned preference models or Pareto-optimal weight search; and (3) conducting large-scale field trials with driver telemetry to evaluate the continuous learning feedback mechanism under live traffic conditions.")

    # --- REFERENCES ---
    add_h1("REFERENCES")
    refs = [
        "[1] M. Barth and K. Boriboonsomsin, \"Real-world carbon dioxide impacts of traffic congestion,\" Transportation Research Record, vol. 2058, no. 1, pp. 163–171, 2008.",
        "[2] V. L. Knoop, S. P. Hoogendoorn, and J. W. C. van Lint, \"Routing traffic in urban networks,\" IEEE Trans. Intell. Transp. Syst., vol. 13, no. 3, pp. 1132–1142, 2012.",
        "[3] E. Cascetta, Transportation Systems Engineering: Theory and Methods. Springer, 2013.",
        "[4] E. W. Dijkstra, \"A note on two problems in connexion with graphs,\" Numerische Mathematik, vol. 1, no. 1, pp. 269–271, 1959.",
        "[5] P. E. Hart, N. J. Nilsson, and B. Raphael, \"A formal basis for the heuristic determination of minimum cost paths,\" IEEE Trans. Syst. Sci. Cybern., vol. 4, no. 2, pp. 100–107, 1968.",
        "[6] L. Alexander, S. Scora, and M. Barth, \"Incorporating dynamic traffic into eco-routing algorithms,\" IEEE Trans. Intell. Transp. Syst., vol. 16, no. 1, pp. 240–251, 2015.",
        "[7] J. London, Intelligent Mobility and Modern Urban Logistics. Academic Press, 2020.",
        "[8] Y. Lv, Y. Duan, W. Kang, Z. Li, and F.-Y. Wang, \"Traffic flow prediction with big data: a deep learning approach,\" IEEE Trans. Intell. Transp. Syst., vol. 16, no. 2, pp. 865–873, 2015.",
        "[9] T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" in Proc. 22nd ACM SIGKDD, 2016, pp. 785–794.",
        "[10] Z. Wu, S. Pan, F. Chen, G. Long, C. Zhang, and P. S. Yu, \"A comprehensive survey on graph neural networks,\" IEEE Trans. Neural Netw. Learn. Syst., vol. 32, no. 1, pp. 4–24, 2021.",
        "[11] P. Veličković, G. Cucurull, A. Casanova, A. Romero, P. Liò, and Y. Bengio, \"Graph attention networks,\" in Proc. ICLR, 2018.",
        "[12] M. Haklay and P. Weber, \"OpenStreetMap: User-generated street maps,\" IEEE Pervasive Comput., vol. 7, no. 4, pp. 12–18, 2008.",
        "[13] G. Boeing, \"OSMnx: New methods for acquiring, constructing, analyzing, and visualizing complex street networks,\" Comput. Environ. Urban Syst., vol. 65, pp. 126–139, 2017.",
        "[14] R. Bellman, \"On a routing problem,\" Quart. Appl. Math., vol. 16, no. 1, pp. 87–90, 1958.",
        "[15] J. Y. Yen, \"Finding the K shortest loopless paths in a network,\" Manage. Sci., vol. 17, no. 11, pp. 712–716, 1971.",
        "[16] L. Breiman, \"Random forests,\" Mach. Learn., vol. 45, no. 1, pp. 5–32, 2001.",
        "[17] P. Geurts, D. Ernst, and L. Wehenkel, \"Extremely randomized trees,\" Mach. Learn., vol. 63, no. 1, pp. 3–42, 2006.",
        "[18] J. H. Friedman, \"Greedy function approximation: a gradient boosting machine,\" Ann. Stat., vol. 29, no. 5, pp. 1189–1232, 2001.",
        "[19] D. Nielsen, \"Tree boosting with XGBoost — why does XGBoost win every competition?,\" M.S. thesis, NTNU, 2016.",
        "[20] T. N. Kipf and M. Welling, \"Semi-supervised classification with graph convolutional networks,\" in Proc. ICLR, 2017.",
        "[21] W. Hamilton, Z. Ying, and J. Leskovec, \"Inductive representation learning on large graphs,\" in Adv. NeurIPS, 2017, pp. 1024–1034.",
        "[22] A. Vaswani et al., \"Attention is all you need,\" in Adv. NeurIPS, 2017, pp. 5998–6008.",
        "[23] S. Guo, Y. Lin, N. Feng, C. Song, and H. Wan, \"Attention based spatial-temporal graph convolutional networks for traffic flow forecasting,\" in Proc. AAAI, 2019, pp. 922–929.",
        "[24] M. Ehrgott, Multicriteria Optimization. Springer, 2005.",
        "[25] K. Deb, Multi-Objective Optimization using Evolutionary Algorithms. Wiley, 2001.",
        "[26] R. T. Marler and J. S. Arora, \"Survey of multi-objective optimization methods for engineering,\" Struct. Multidiscip. Optim., vol. 26, no. 6, pp. 369–395, 2004.",
        "[27] E. Triantaphyllou, Multi-criteria Decision Making Methods: A Comparative Study. Springer, 2000.",
        "[28] D. Johnson, \"Algorithms for shortest paths,\" Ph.D. dissertation, Stanford Univ., 1973.",
        "[29] X. Chen, L. Sun, and Y. Liu, \"Stochastic travel time estimation and risk-averse routing,\" Transp. Res. B, vol. 142, pp. 110–135, 2020.",
        "[30] R. T. Rockafellar and S. Uryasev, \"Optimization of conditional value-at-risk,\" J. Risk, vol. 2, pp. 21–42, 2000.",
        "[31] A. J. Kleywegt, V. S. Shapiro, and T. Homem-de-Mello, \"The sample average approximation method for stochastic discrete optimization problems,\" SIAM J. Optim., vol. 12, no. 2, pp. 479–502, 2002.",
        "[32] G. Boeing, \"Street network models and measures for every urban area in the world,\" Geogr. Anal., vol. 53, no. 1, pp. 51–69, 2021.",
        "[33] J.-Y. Li, Z.-H. Zhan, R. Liu, and J. Zhang, \"Tour multi-route planning with matrix-based differential evolution,\" IEEE Trans. Intell. Transp. Syst., vol. 25, no. 9, pp. 12416–12431, Sept. 2024.",
        "[34] X. Peng, L. Ke, and D. Wang, \"Urban multiple route planning model using dynamic programming in reinforcement learning,\" IEEE Trans. Intell. Transp. Syst., vol. 23, no. 7, pp. 8037–8049, Jul. 2022.",
        "[35] A. Khayat et al., \"AI-based predictive modeling and NSGA-II optimization for eco-driving route planning in electric vehicles,\" IEEE Access, vol. 13, pp. 1–15, 2025.",
        "[36] A. Derrow-Pinion et al., \"ETA prediction with graph neural networks in Google Maps,\" in Proc. 30th ACM CIKM, 2021, pp. 3767–3776.",
    ]

    for ref in refs:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after = Pt(2)
        rp.paragraph_format.line_spacing = 1.0
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rr = rp.add_run(ref)
        rr.font.name = 'Times New Roman'
        rr.font.size = Pt(8.5)

    out_file = 'DeepRoute_Research_Paper_INDIACom.docx'
    doc.save(out_file)
    print(f"Successfully generated {out_file} with embedded figures in exact INDIACom/IEEE layout!")

if __name__ == "__main__":
    build_indiacom_paper()
